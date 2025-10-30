from fastapi import FastAPI, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
import os
import uuid

app = FastAPI()

class Chapter(BaseModel):
    title: str
    content: str

class ResearchData(BaseModel):
    title: str
    students: list
    supervisor: str
    college: str
    chapters: list[Chapter]

@app.post("/generate")
def generate_research(data: ResearchData):
    # إنشاء ملف Word جديد
    doc = Document()

    # الغلاف
    doc.add_heading(data.title, 0)
    doc.add_paragraph(f"Supervised by: {data.supervisor}")
    doc.add_paragraph(f"Students: {', '.join(data.students)}")
    doc.add_paragraph(f"College: {data.college}")
    doc.add_page_break()

    # إضافة الفصول
    for chapter in data.chapters:
        doc.add_heading(chapter.title, level=1)
        doc.add_paragraph(chapter.content)
        doc.add_page_break()

    # حفظ الملف
    filename = f"{uuid.uuid4().hex[:8]}{data.title.replace(' ', '')}.docx"
    filepath = f"/tmp/{filename}"
    doc.save(filepath)

    return FileResponse(filepath, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)
    
    @app.get("/")
def home():
    return {"message": "FastAPI is running successfully!"}
