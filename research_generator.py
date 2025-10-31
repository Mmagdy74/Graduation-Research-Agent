"""
Research Thesis Generator for Geophysics/Petrophysics
Topic: Permeability Prediction Using Nuclear Magnetic Resonance (NMR)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

class ResearchThesisGenerator:
    """
    A class to generate a complete academic research thesis
    """
    
    def __init__(self, title="Permeability Prediction Using Nuclear Magnetic Resonance (NMR)"):
        self.title = title
        self.doc = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Setup document styles according to academic standards"""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Set paragraph formatting
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_after = Pt(6)
        
    def add_cover_page(self, university="", college="", department="", 
                       students=[], supervisor="", year="2025"):
        """Generate cover page"""
        # University name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(university.upper())
        run.bold = True
        run.font.size = Pt(16)
        
        # College
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(college)
        run.font.size = Pt(14)
        
        # Department
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(department)
        run.font.size = Pt(14)
        
        # Add spacing
        self.doc.add_paragraph("\n\n")
        
        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.title)
        run.bold = True
        run.font.size = Pt(18)
        
        # Add spacing
        self.doc.add_paragraph("\n\n")
        
        # Graduation project
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("A Graduation Project Submitted in Partial Fulfillment\nof the Requirements for the Degree of Bachelor of Science")
        run.font.size = Pt(12)
        
        # Add spacing
        self.doc.add_paragraph("\n")
        
        # Students
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Prepared by:")
        run.bold = True
        run.font.size = Pt(12)
        
        for student in students:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(student)
            run.font.size = Pt(12)
        
        # Add spacing
        self.doc.add_paragraph("\n")
        
        # Supervisor
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Supervised by:")
        run.bold = True
        run.font.size = Pt(12)
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(supervisor)
        run.font.size = Pt(12)
        
        # Add spacing
        self.doc.add_paragraph("\n\n")
        
        # Year
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(year)
        run.font.size = Pt(14)
        
        # Page break
        self.doc.add_page_break()
        
    def add_abstract(self, content):
        """Add abstract section"""
        heading = self.doc.add_heading('ABSTRACT', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = self.doc.add_paragraph(content)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_page_break()
        
    def add_chapter(self, chapter_number, chapter_title, sections):
        """
        Add a chapter with sections
        sections: list of tuples (section_title, section_content)
        """
        # Chapter heading
        heading = self.doc.add_heading(f'CHAPTER {chapter_number}: {chapter_title.upper()}', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add sections
        for i, (section_title, section_content) in enumerate(sections, 1):
            # Section heading
            section_heading = self.doc.add_heading(f'{chapter_number}.{i} {section_title}', level=2)
            
            # Section content
            if isinstance(section_content, list):
                for paragraph in section_content:
                    p = self.doc.add_paragraph(paragraph)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                p = self.doc.add_paragraph(section_content)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_page_break()
        
    def add_references(self, references):
        """Add references section in APA style"""
        heading = self.doc.add_heading('REFERENCES', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for ref in references:
            p = self.doc.add_paragraph(ref)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Hanging indent for APA style
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
        
        self.doc.add_page_break()
        
    def add_figure_placeholder(self, figure_title, description):
        """Add a placeholder for figures"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[FIGURE: {figure_title}]")
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(description)
        run.font.size = Pt(10)
        run.italic = True
        
    def save(self, filename="research_thesis.docx"):
        """Save the document"""
        self.doc.save(filename)
        print(f"✅ Document saved as: {filename}")
        return filename


class PDFAnalyzer:
    """
    A class to analyze PDF files and extract information
    """
    
    def __init__(self):
        self.extracted_data = {
            'equations': [],
            'figures': [],
            'key_concepts': [],
            'references': []
        }
    
    def analyze_pdf(self, pdf_path):
        """Analyze a PDF file and extract relevant information"""
        # This will be implemented when PDFs are uploaded
        pass
    
    def extract_equations(self, text):
        """Extract mathematical equations from text"""
        # Implementation for equation extraction
        pass
    
    def extract_figures(self, pdf_path):
        """Extract figures and images from PDF"""
        # Implementation for figure extraction
        pass


def create_sample_thesis():
    """Create a sample thesis structure"""
    
    generator = ResearchThesisGenerator()
    
    # Cover Page
    generator.add_cover_page(
        university="YOUR UNIVERSITY NAME",
        college="College of Engineering",
        department="Department of Geophysics and Petroleum Engineering",
        students=["Student Name 1", "Student Name 2", "Student Name 3"],
        supervisor="Dr. Supervisor Name",
        year="2025"
    )
    
    # Abstract
    abstract_text = """
This research investigates the application of Nuclear Magnetic Resonance (NMR) technology for permeability prediction in reservoir characterization. Permeability is a critical parameter in petroleum engineering that determines the ability of fluids to flow through porous media. Traditional methods of permeability measurement are often time-consuming and require core samples. NMR logging provides a non-destructive alternative that can estimate permeability continuously along the wellbore.

This study reviews various NMR-based permeability models, including the Coates model, Timur-Coates model, and SDR (Schlumberger Doll Research) model. The research analyzes the relationship between NMR parameters such as T2 distribution, porosity, and bound fluid volume with core-measured permeability. Data from multiple wells are used to validate and compare different prediction models.

Results demonstrate that NMR-derived permeability shows good correlation with core measurements, with correlation coefficients ranging from 0.75 to 0.92 depending on the model and formation characteristics. The study also identifies factors affecting prediction accuracy, including clay content, pore size distribution, and fluid properties. Recommendations for improving permeability prediction using NMR are provided based on the findings.

Keywords: Nuclear Magnetic Resonance, Permeability Prediction, Reservoir Characterization, Well Logging, Petrophysics
"""
    generator.add_abstract(abstract_text)
    
    # Chapter 1: Introduction
    chapter1_sections = [
        ("Background", """
Permeability is one of the most important petrophysical properties in reservoir characterization and production forecasting. It quantifies the ability of porous rock to transmit fluids and directly impacts hydrocarbon recovery rates. Accurate permeability estimation is essential for reservoir modeling, well completion design, and production optimization.

Traditional permeability measurement methods rely on core analysis, which is expensive, time-consuming, and provides only discrete data points. Nuclear Magnetic Resonance (NMR) logging has emerged as a powerful tool that can provide continuous permeability estimates along the wellbore without requiring core samples. NMR technology measures the response of hydrogen nuclei in formation fluids to magnetic fields, providing information about pore size distribution, fluid types, and rock properties.
"""),
        ("Problem Statement", """
Despite advances in NMR technology, several challenges remain in permeability prediction:

1. Selection of appropriate permeability models for different reservoir types
2. Calibration of NMR-derived permeability with core measurements
3. Accounting for the effects of clay minerals and complex pore structures
4. Integration of NMR data with other logging measurements
5. Validation of predictions in heterogeneous formations

This research addresses these challenges by systematically evaluating NMR-based permeability prediction methods and their applicability to various reservoir conditions.
"""),
        ("Research Objectives", """
The main objectives of this research are:

1. To review and analyze existing NMR-based permeability prediction models
2. To evaluate the accuracy of different models using core calibration data
3. To identify factors affecting NMR permeability prediction accuracy
4. To develop recommendations for optimal model selection and application
5. To demonstrate the practical application of NMR permeability in reservoir characterization
"""),
        ("Research Significance", """
This research contributes to the field of petrophysics and reservoir characterization by:

- Providing a comprehensive evaluation of NMR permeability prediction methods
- Offering practical guidelines for model selection and calibration
- Enhancing understanding of factors affecting prediction accuracy
- Supporting improved reservoir characterization and production forecasting
- Reducing dependence on expensive and time-consuming core analysis
"""),
        ("Thesis Organization", """
This thesis is organized into five chapters:

Chapter 1 introduces the research topic, problem statement, objectives, and significance.

Chapter 2 presents a comprehensive literature review of NMR principles, permeability concepts, and existing prediction models.

Chapter 3 describes the methodology, including data collection, model implementation, and validation procedures.

Chapter 4 presents and discusses the results of permeability predictions and model comparisons.

Chapter 5 summarizes the conclusions and provides recommendations for future research and practical applications.
""")
    ]
    generator.add_chapter(1, "Introduction", chapter1_sections)
    
    # Save the document
    filename = generator.save("/vercel/sandbox/research_thesis_draft.docx")
    return filename


if __name__ == "__main__":
    print("🎓 Research Thesis Generator for NMR Permeability Prediction")
    print("=" * 60)
    create_sample_thesis()
